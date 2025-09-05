"""LlamaIndex-based DOCX Parser."""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from .base_parser import BaseLlamaIndexParser, ParserFactory

logger = logging.getLogger(__name__)

try:
    from llama_index.readers.file import DocxReader
    LLAMA_DOCX_AVAILABLE = True
except ImportError:
    LLAMA_DOCX_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class LlamaIndexDocxParser(BaseLlamaIndexParser):
    """LlamaIndex-based parser for DOCX files."""
    
    def __init__(self, name: str = "LlamaIndexDocxParser", config: Optional[Dict[str, Any]] = None):
        """
        Initialize LlamaIndex DOCX parser.
        
        Args:
            name: Parser name
            config: Parser configuration
        """
        super().__init__(name=name, config=config or {})
        
        # DOCX-specific configuration
        self.extract_metadata = self.config.get("extract_metadata", True)
        self.extract_headers_footers = self.config.get("extract_headers_footers", True)
        self.extract_comments = self.config.get("extract_comments", False)
        self.extract_tables = self.config.get("extract_tables", True)
        self.extract_images = self.config.get("extract_images", False)
        self.preserve_formatting = self.config.get("preserve_formatting", False)
        
        if not LLAMA_DOCX_AVAILABLE and not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "Either llama-index DocxReader or python-docx is required. "
                "Install with: pip install llama-index or pip install python-docx"
            )
    
    def _get_reader(self):
        """Get the appropriate LlamaIndex reader for DOCX files."""
        if LLAMA_DOCX_AVAILABLE:
            return DocxReader()
        else:
            # Will use manual parsing with python-docx
            return None
    
    def parse(self, file_path: str, **kwargs) -> "ProcessingResult":
        """
        Parse a DOCX file using LlamaIndex.
        
        Args:
            file_path: Path to the DOCX file
            **kwargs: Additional parsing options
            
        Returns:
            ProcessingResult containing parsed documents
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            from core.base import ProcessingResult
            return ProcessingResult(
                documents=[],
                errors=[{"error": f"File not found: {file_path}", "source": str(file_path)}]
            )
        
        try:
            if self._reader:
                return self._parse_with_llamaindex(file_path, **kwargs)
            else:
                return self._parse_with_python_docx(file_path, **kwargs)
        
        except Exception as e:
            from core.base import ProcessingResult
            return ProcessingResult(
                documents=[],
                errors=[{
                    "error": f"Failed to parse DOCX file: {str(e)}",
                    "source": str(file_path)
                }]
            )
    
    def _parse_with_llamaindex(self, file_path: Path, **kwargs) -> "ProcessingResult":
        """Parse DOCX using LlamaIndex DocxReader."""
        try:
            documents = self._reader.load_data(file=str(file_path))
            
            if not documents:
                from core.base import ProcessingResult
                return ProcessingResult(
                    documents=[],
                    errors=[{"error": "No content extracted by LlamaIndex", "source": str(file_path)}]
                )
            
            return self._process_llamaindex_documents(documents, file_path)
            
        except Exception as e:
            logger.warning(f"LlamaIndex DOCX reader failed: {e}, falling back to python-docx")
            if PYTHON_DOCX_AVAILABLE:
                return self._parse_with_python_docx(file_path, **kwargs)
            else:
                from core.base import ProcessingResult
                return ProcessingResult(
                    documents=[],
                    errors=[{
                        "error": f"LlamaIndex parsing failed and no fallback available: {str(e)}",
                        "source": str(file_path)
                    }]
                )
    
    def _parse_with_python_docx(self, file_path: Path, **kwargs) -> "ProcessingResult":
        """Parse DOCX using python-docx as fallback."""
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract content
            content_parts = []
            
            # Extract main document text
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extract tables if enabled
            if self.extract_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        content_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                from core.base import ProcessingResult
                return ProcessingResult(
                    documents=[],
                    errors=[{"error": "No text content found in DOCX", "source": str(file_path)}]
                )
            
            # Extract metadata
            metadata = self._extract_docx_metadata(doc, file_path)
            
            return self._create_documents_from_content(content, metadata, file_path)
            
        except Exception as e:
            from core.base import ProcessingResult
            return ProcessingResult(
                documents=[],
                errors=[{
                    "error": f"python-docx parsing failed: {str(e)}",
                    "source": str(file_path)
                }]
            )
    
    def _process_llamaindex_documents(self, documents, file_path: Path) -> "ProcessingResult":
        """Process LlamaIndex documents into our format."""
        from core.base import ProcessingResult, Document
        from utils.hash_utils import generate_document_metadata, generate_chunk_metadata
        
        result_documents = []
        errors = []
        
        # Combine all content if multiple documents
        if len(documents) > 1:
            combined_content = "\n\n".join(doc.text for doc in documents if doc.text)
        else:
            combined_content = documents[0].text if documents else ""
        
        if not combined_content.strip():
            return ProcessingResult(
                documents=[],
                errors=[{"error": "No content extracted", "source": str(file_path)}]
            )
        
        # Generate metadata
        base_metadata = generate_document_metadata(str(file_path), combined_content)
        base_metadata.update({
            "parser_type": self.name,
            "reader_type": "llamaindex_docx"
        })
        
        # Add LlamaIndex metadata from first document
        if documents[0].metadata:
            base_metadata.update(documents[0].metadata)
        
        # Handle chunking
        if self._node_parser and self.chunk_size:
            chunked_docs = self._create_chunked_documents(
                combined_content, base_metadata, documents[0]
            )
            result_documents.extend(chunked_docs)
        else:
            document_id = f"doc_{base_metadata['document_hash'][:12]}_full"
            doc = Document(
                content=combined_content,
                metadata=base_metadata,
                id=document_id,
                source=str(file_path)
            )
            result_documents.append(doc)
        
        return ProcessingResult(
            documents=result_documents,
            errors=errors,
            metrics={
                "total_documents": len(result_documents),
                "total_errors": len(errors),
                "file_processed": str(file_path),
                "parser_type": self.name
            }
        )
    
    def _create_documents_from_content(self, content: str, metadata: Dict[str, Any], file_path: Path) -> "ProcessingResult":
        """Create documents from extracted content."""
        from core.base import ProcessingResult, Document
        from utils.hash_utils import generate_document_metadata, generate_chunk_metadata
        
        # Generate metadata
        base_metadata = generate_document_metadata(str(file_path), content)
        base_metadata.update(metadata)
        
        result_documents = []
        
        # Handle chunking
        if self._node_parser and self.chunk_size:
            # Create mock LlamaIndex document for chunking
            from llama_index.core.schema import Document as LlamaDocument
            llama_doc = LlamaDocument(text=content, metadata=base_metadata)
            
            chunked_docs = self._create_chunked_documents(
                content, base_metadata, llama_doc
            )
            result_documents.extend(chunked_docs)
        else:
            document_id = f"doc_{base_metadata['document_hash'][:12]}_full"
            doc = Document(
                content=content,
                metadata=base_metadata,
                id=document_id,
                source=str(file_path)
            )
            result_documents.append(doc)
        
        return ProcessingResult(
            documents=result_documents,
            errors=[],
            metrics={
                "total_documents": len(result_documents),
                "file_processed": str(file_path),
                "parser_type": self.name
            }
        )
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table."""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
    
    def _extract_docx_metadata(self, doc, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX document."""
        metadata = {
            "parser_type": self.name,
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "reader_type": "python_docx"
        }
        
        try:
            # Core properties
            core_props = doc.core_properties
            if core_props:
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "comments": core_props.comments or "",
                    "category": core_props.category or "",
                    "created": core_props.created.isoformat() if core_props.created else "",
                    "modified": core_props.modified.isoformat() if core_props.modified else "",
                    "last_modified_by": core_props.last_modified_by or "",
                    "revision": str(core_props.revision) if core_props.revision else ""
                })
            
            # Document statistics
            if hasattr(doc, 'paragraphs'):
                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                metadata["paragraph_count"] = paragraph_count
            
            if hasattr(doc, 'tables') and self.extract_tables:
                metadata["table_count"] = len(doc.tables)
            
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
        
        return metadata
    
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the given file."""
        return Path(file_path).suffix.lower() in self.get_supported_extensions()
    
    @staticmethod
    def can_parse_mime_type(mime_type: str) -> bool:
        """Check if this parser can handle the given MIME type."""
        return mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    
    @staticmethod
    def get_supported_extensions() -> List[str]:
        """Get list of supported file extensions."""
        return ['.docx', '.doc']
    
    @staticmethod
    def get_description() -> str:
        """Get parser description."""
        return "LlamaIndex-based parser for DOCX files"


# Register the parser
ParserFactory.register_parser("LlamaIndexDocxParser", LlamaIndexDocxParser)